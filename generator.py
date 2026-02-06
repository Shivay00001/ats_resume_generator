
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from constants import (
    FONT_NAME, FONT_SIZE_BODY, FONT_SIZE_HEADING, FONT_SIZE_NAME,
    MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT, MARGIN_RIGHT
)

def generate_resume(data, output_path, export_format='docx'):
    """
    Dispatcher for resume generation.
    """
    if export_format == 'docx':
        return params_to_docx(data, output_path)
    elif export_format == 'pdf':
        return params_to_pdf(data, output_path)
    elif export_format == 'txt':
        return params_to_txt(data, output_path)
    else:
        raise ValueError(f"Unsupported format: {export_format}")

def params_to_docx(data, output_path):
    doc = Document()
    
    # Setup Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(MARGIN_TOP)
        section.bottom_margin = Inches(MARGIN_BOTTOM)
        section.left_margin = Inches(MARGIN_LEFT)
        section.right_margin = Inches(MARGIN_RIGHT)
    
    # Setup Default Font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE_BODY)

    # --- Header (Contact Info) ---
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = name_paragraph.add_run(data.get('full_name', '').upper())
    run.bold = True
    run.font.size = Pt(FONT_SIZE_NAME)
    run.font.name = FONT_NAME

    contact_lines = []
    if data.get('city') or data.get('country'):
        contact_lines.append(f"{data.get('city', '')}, {data.get('country', '')}")
    if data.get('phone'):
        contact_lines.append(data.get('phone'))
    if data.get('email'):
        contact_lines.append(data.get('email'))
    if data.get('linkedin'):
        contact_lines.append(data.get('linkedin'))
    if data.get('github'):
        contact_lines.append(data.get('github'))
    
    contact_paragraph = doc.add_paragraph(" | ".join([c for c in contact_lines if c]))
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contact_paragraph.paragraph_format.space_after = Pt(12)

    # --- Helper for Heading ---
    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(FONT_SIZE_HEADING)
        run.font.name = FONT_NAME
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        # Add a simple bottom border if possible, or just a line divider? 
        # Word borders are tricky with python-docx without XML hacking.
        # Staying simple (No visual lines to avoid graphics issues is safer for strict ATS, 
        # though border-bottom is usually fine. We'll stick to bold caps).

    # --- Professional Summary ---
    if data.get('summary'):
        add_heading("Professional Summary")
        p = doc.add_paragraph(data.get('summary'))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Skills ---
    if data.get('skills'):
        add_heading("Skills")
        # Just comma separated as requested in some ATS formats to save space, 
        # or bullet points? Prompt says "Skills (comma-separated)" in form.
        # Let's display them as a clean paragraph or list. Paragraph is safer for density.
        p = doc.add_paragraph(data.get('skills'))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Work Experience ---
    experiences = data.get('experience', [])
    if experiences:
        add_heading("Work Experience")
        for exp in experiences:
            # Header: Role | Company | Location | Dates 
            # (Or variations. Standard is typically: Company on left, Date on right.
            # But plain text usually easier if linear or pipe separated)
            
            # Line 1: Company -- Location
            p_comp = doc.add_paragraph()
            r_comp = p_comp.add_run(f"{exp.get('company', '')}")
            r_comp.bold = True
            if exp.get('location'):
                p_comp.add_run(f" | {exp.get('location', '')}")

            # Line 2: Title -- Dates
            p_role = doc.add_paragraph()
            r_role = p_role.add_run(f"{exp.get('title', '')}")
            r_role.italic = True
            if exp.get('start_date') or exp.get('end_date'):
                p_role.add_run(f" | {exp.get('start_date', '')} - {exp.get('end_date', '')}")
            
            # Bullets
            bullets = exp.get('responsibilities', '').split('\n')
            for b in bullets:
                if b.strip():
                    p_bull = doc.add_paragraph(b.strip(), style='List Bullet')
                    p_bull.paragraph_format.space_after = Pt(2)

    # --- Projects ---
    projects = data.get('projects', [])
    if projects:
        add_heading("Projects")
        for proj in projects:
            p_proj = doc.add_paragraph()
            r_proj = p_proj.add_run(f"{proj.get('name', '')}")
            r_proj.bold = True
            
            bullets = proj.get('description', '').split('\n')
            for b in bullets:
                if b.strip():
                    p_bull = doc.add_paragraph(b.strip(), style='List Bullet')
                    p_bull.paragraph_format.space_after = Pt(2)

    # --- Education ---
    education = data.get('education', [])
    if education:
        add_heading("Education")
        for edu in education:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.get('institution', '')}")
            run.bold = True
            
            p2 = doc.add_paragraph()
            p2.add_run(f"{edu.get('degree', '')}")
            if edu.get('year'):
                p2.add_run(f" | {edu.get('year', '')}")

    # --- Certifications ---
    certifications = data.get('certifications', [])
    if certifications:
        add_heading("Certifications")
        for cert in certifications:
             doc.add_paragraph(cert.strip(), style='List Bullet')

    doc.save(output_path)
    return output_path

def params_to_txt(data, output_path):
    lines = []
    
    # Header
    lines.append(data.get('full_name', '').upper())
    contact = []
    if data.get('phone'): contact.append(data.get('phone'))
    if data.get('email'): contact.append(data.get('email'))
    if data.get('linkedin'): contact.append(data.get('linkedin'))
    if data.get('city'): contact.append(f"{data.get('city')}, {data.get('country', '')}")
    if data.get('github'): contact.append(data.get('github'))
    lines.append(" | ".join(contact))
    lines.append("")
    
    def add_sec(title, content_fn):
        lines.append(title.upper())
        lines.append("-" * len(title))
        content_fn()
        lines.append("")

    if data.get('summary'):
        def summary_content():
            lines.append(data.get('summary'))
        add_sec("Professional Summary", summary_content)
        
    if data.get('skills'):
        def skills_content():
            lines.append(data.get('skills'))
        add_sec("Skills", skills_content)
        
    if data.get('experience'):
        def exp_content():
            for exp in data.get('experience'):
                lines.append(f"{exp.get('company')} | {exp.get('location', '')}")
                lines.append(f"{exp.get('title')} | {exp.get('start_date', '')} - {exp.get('end_date', '')}")
                for line in exp.get('responsibilities', '').split('\n'):
                    if line.strip():
                        lines.append(f"- {line.strip()}")
                lines.append("")
        add_sec("Work Experience", exp_content)

    if data.get('projects'):
        def proj_content():
            for proj in data.get('projects'):
                lines.append(f"{proj.get('name')}")
                for line in proj.get('description', '').split('\n'):
                    if line.strip():
                        lines.append(f"- {line.strip()}")
                lines.append("")
        add_sec("Projects", proj_content)

    if data.get('education'):
        def edu_content():
            for edu in data.get('education'):
                lines.append(f"{edu.get('institution')}")
                lines.append(f"{edu.get('degree')} | {edu.get('year', '')}")
                lines.append("")
        add_sec("Education", edu_content)
        
    if data.get('certifications'):
        def cert_content():
            for cert in data.get('certifications'):
                if cert.strip():
                    lines.append(f"- {cert.strip()}")
        add_sec("Certifications", cert_content)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines))
    return output_path

def params_to_pdf(data, output_path):
    # Plain PDF using ReportLab
    doc = SimpleDocTemplate(
        output_path,
        pagesize=LETTER,
        rightMargin=72, leftMargin=72,
        topMargin=72, bottomMargin=72
    )
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    style_name = ParagraphStyle('Name', parent=styles['Normal'], fontSize=14, spaceAfter=6, fontName='Helvetica-Bold')
    style_heading = ParagraphStyle('Heading', parent=styles['Normal'], fontSize=12, spaceAfter=6, spaceBefore=12, fontName='Helvetica-Bold')
    style_body = ParagraphStyle('Body', parent=styles['Normal'], fontSize=11, spaceAfter=2, fontName='Helvetica')
    style_bullet = ParagraphStyle('Bullet', parent=styles['Normal'], fontSize=11, leftIndent=20, spaceAfter=2, fontName='Helvetica', bulletText='•')

    # Header
    story.append(Paragraph(data.get('full_name', '').upper(), style_name))
    
    contact_info = []
    if data.get('email'): contact_info.append(data.get('email'))
    if data.get('phone'): contact_info.append(data.get('phone'))
    if data.get('linkedin'): contact_info.append(data.get('linkedin'))
    story.append(Paragraph(" | ".join(contact_info), style_body))
    story.append(Spacer(1, 12))

    # Summary
    if data.get('summary'):
        story.append(Paragraph("PROFESSIONAL SUMMARY", style_heading))
        story.append(Paragraph(data.get('summary', ''), style_body))

    # Skills
    if data.get('skills'):
        story.append(Paragraph("SKILLS", style_heading))
        story.append(Paragraph(data.get('skills', ''), style_body))

    # Experience
    if data.get('experience'):
        story.append(Paragraph("WORK EXPERIENCE", style_heading))
        for exp in data.get('experience'):
            header = f"<b>{exp.get('company')}</b> | {exp.get('location')} <br/><i>{exp.get('title')}</i> | {exp.get('start_date')} - {exp.get('end_date')}"
            story.append(Paragraph(header, style_body))
            for line in exp.get('responsibilities', '').split('\n'):
                if line.strip():
                    story.append(Paragraph(line.strip(), style_bullet))
            story.append(Spacer(1, 6))

    # Projects
    if data.get('projects'):
        story.append(Paragraph("PROJECTS", style_heading))
        for proj in data.get('projects'):
            story.append(Paragraph(f"<b>{proj.get('name')}</b>", style_body))
            for line in proj.get('description', '').split('\n'):
                if line.strip():
                    story.append(Paragraph(line.strip(), style_bullet))
            story.append(Spacer(1, 6))

    # Education
    if data.get('education'):
        story.append(Paragraph("EDUCATION", style_heading))
        for edu in data.get('education'):
            story.append(Paragraph(f"<b>{edu.get('institution')}</b>", style_body))
            story.append(Paragraph(f"{edu.get('degree')} | {edu.get('year')}", style_body))
            story.append(Spacer(1, 6))
            
    # Certifications
    if data.get('certifications'):
        story.append(Paragraph("CERTIFICATIONS", style_heading))
        for cert in data.get('certifications'):
            if cert.strip():
                story.append(Paragraph(cert.strip(), style_bullet))

    doc.build(story)
    return output_path
